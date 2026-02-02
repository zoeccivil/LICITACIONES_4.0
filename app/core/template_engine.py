"""
Template Engine Module - Motor de plantillas y generación de documentos

Este módulo maneja la generación automática de documentos (cartas, oficios, etc.)
utilizando plantillas con variables dinámicas.

Soporta:
- Plantillas DOCX (python-docx)
- Plantillas HTML (Jinja2)
- Variables dinámicas: {{variable}}
- Guardado en Firebase Storage (opcional)
"""
from __future__ import annotations

import os
import datetime
from typing import Any, Dict, List, Optional
from pathlib import Path

# Constant for placeholder pattern
PLACEHOLDER_PATTERN = "{{{{{key}}}}}"

try:
    from docx import Document
    from docx.shared import Pt, Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from jinja2 import Environment, FileSystemLoader, Template
    JINJA2_AVAILABLE = True
except ImportError:
    JINJA2_AVAILABLE = False


class TemplateEngine:
    """Motor de generación de documentos desde plantillas."""

    def __init__(self, templates_dir: Optional[str] = None):
        """
        Inicializa el motor de plantillas.

        Args:
            templates_dir: Directorio donde se almacenan las plantillas
        """
        self.templates_dir = templates_dir or os.path.join(os.getcwd(), "templates")
        
        # Crear directorio si no existe
        Path(self.templates_dir).mkdir(parents=True, exist_ok=True)
        
        # Inicializar Jinja2 si está disponible
        if JINJA2_AVAILABLE:
            self.jinja_env = Environment(loader=FileSystemLoader(self.templates_dir))
        else:
            self.jinja_env = None

    def list_templates(self) -> List[str]:
        """
        Lista todas las plantillas disponibles.

        Returns:
            Lista de nombres de archivos de plantillas
        """
        templates_path = Path(self.templates_dir)
        if not templates_path.exists():
            return []
        
        templates = []
        for file in templates_path.glob("*"):
            if file.suffix in [".docx", ".html", ".txt"]:
                templates.append(file.name)
        
        return sorted(templates)

    def generate_from_docx_template(
        self,
        template_name: str,
        variables: Dict[str, Any],
        output_path: str
    ) -> bool:
        """
        Genera un documento DOCX desde una plantilla.

        Args:
            template_name: Nombre del archivo de plantilla
            variables: Diccionario con variables a reemplazar
            output_path: Ruta donde guardar el documento generado

        Returns:
            True si se generó exitosamente
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx no está instalado. Instálalo con: pip install python-docx")

        template_path = os.path.join(self.templates_dir, template_name)
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Plantilla no encontrada: {template_path}")

        try:
            # Cargar plantilla
            doc = Document(template_path)
            
            # Reemplazar variables en párrafos
            for paragraph in doc.paragraphs:
                for key, value in variables.items():
                    placeholder = PLACEHOLDER_PATTERN.format(key=key)
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, str(value))
            
            # Reemplazar variables en tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in variables.items():
                            placeholder = PLACEHOLDER_PATTERN.format(key=key)
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, str(value))
            
            # Guardar documento
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Error al generar documento DOCX: {e}")
            return False

    def generate_from_html_template(
        self,
        template_name: str,
        variables: Dict[str, Any],
        output_path: str
    ) -> bool:
        """
        Genera un documento HTML desde una plantilla Jinja2.

        Args:
            template_name: Nombre del archivo de plantilla
            variables: Diccionario con variables a reemplazar
            output_path: Ruta donde guardar el documento generado

        Returns:
            True si se generó exitosamente
        """
        if not JINJA2_AVAILABLE:
            raise ImportError("Jinja2 no está instalado. Instálalo con: pip install Jinja2")

        try:
            template = self.jinja_env.get_template(template_name)
            rendered = template.render(**variables)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(rendered)
            
            return True
            
        except Exception as e:
            print(f"Error al generar documento HTML: {e}")
            return False

    def generate_from_string_template(
        self,
        template_string: str,
        variables: Dict[str, Any]
    ) -> str:
        """
        Genera texto desde una plantilla en string.

        Args:
            template_string: Texto de la plantilla con placeholders {{variable}}
            variables: Diccionario con variables a reemplazar

        Returns:
            Texto con variables reemplazadas
        """
        result = template_string
        for key, value in variables.items():
            placeholder = PLACEHOLDER_PATTERN.format(key=key)
            result = result.replace(placeholder, str(value))
        
        return result

    def create_simple_docx(
        self,
        output_path: str,
        titulo: str,
        contenido: str,
        pie_pagina: str = ""
    ) -> bool:
        """
        Crea un documento DOCX simple sin plantilla.

        Args:
            output_path: Ruta donde guardar el documento
            titulo: Título del documento
            contenido: Contenido principal
            pie_pagina: Texto del pie de página (opcional)

        Returns:
            True si se creó exitosamente
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx no está instalado. Instálalo con: pip install python-docx")

        try:
            doc = Document()
            
            # Título
            heading = doc.add_heading(titulo, level=1)
            
            # Contenido (puede contener saltos de línea)
            for linea in contenido.split('\n'):
                doc.add_paragraph(linea)
            
            # Pie de página
            if pie_pagina:
                doc.add_paragraph()
                doc.add_paragraph(pie_pagina)
            
            # Guardar
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Error al crear documento simple: {e}")
            return False

    def generate_carta_oferta(
        self,
        licitacion_data: Dict[str, Any],
        empresa_data: Dict[str, Any],
        output_path: str
    ) -> bool:
        """
        Genera una carta de oferta para una licitación.

        Args:
            licitacion_data: Datos de la licitación
            empresa_data: Datos de la empresa
            output_path: Ruta donde guardar el documento

        Returns:
            True si se generó exitosamente
        """
        variables = {
            "fecha": datetime.date.today().strftime("%d/%m/%Y"),
            "institucion": licitacion_data.get("institucion", ""),
            "numero_proceso": licitacion_data.get("numero_proceso", ""),
            "nombre_proceso": licitacion_data.get("nombre_proceso", ""),
            "razon_social": empresa_data.get("nombre", ""),
            "rnc": empresa_data.get("rnc", ""),
            "direccion": empresa_data.get("direccion", ""),
            "telefono": empresa_data.get("telefono", ""),
            "email": empresa_data.get("email", ""),
        }

        # Buscar plantilla de carta de oferta
        template_name = "carta_oferta.docx"
        template_path = os.path.join(self.templates_dir, template_name)
        
        if os.path.exists(template_path):
            return self.generate_from_docx_template(template_name, variables, output_path)
        else:
            # Crear carta simple si no hay plantilla
            titulo = f"Carta de Oferta - {variables['numero_proceso']}"
            
            contenido = f"""
Estimados señores de {variables['institucion']}:

Por medio de la presente, {variables['razon_social']}, RNC: {variables['rnc']}, 
presenta su oferta formal para el proceso de licitación:

Número de Proceso: {variables['numero_proceso']}
Nombre: {variables['nombre_proceso']}

Quedamos a su disposición para cualquier aclaración.

Atentamente,
{variables['razon_social']}
{variables['direccion']}
Tel: {variables['telefono']}
Email: {variables['email']}
            """.strip()
            
            pie_pagina = f"Fecha: {variables['fecha']}"
            
            return self.create_simple_docx(output_path, titulo, contenido, pie_pagina)

    def get_available_variables(self) -> Dict[str, str]:
        """
        Retorna las variables disponibles para usar en plantillas.

        Returns:
            Diccionario con nombre de variable y descripción
        """
        return {
            "fecha": "Fecha actual (DD/MM/YYYY)",
            "institucion": "Nombre de la institución",
            "numero_proceso": "Número del proceso de licitación",
            "nombre_proceso": "Nombre del proceso",
            "razon_social": "Razón social de la empresa",
            "rnc": "RNC de la empresa",
            "direccion": "Dirección de la empresa",
            "telefono": "Teléfono de contacto",
            "email": "Email de contacto",
            "monto": "Monto ofertado",
            "proyecto": "Nombre del proyecto",
            "lote": "Número de lote",
            "responsable": "Nombre del responsable",
        }
